from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Style helpers ─────────────────────────────────────────────────────────────
BRAND   = RGBColor(0x1F, 0x77, 0xB4)
DARK    = RGBColor(0x22, 0x22, 0x22)
GREY    = RGBColor(0x55, 0x55, 0x55)

def set_font(run, size=11, bold=False, italic=False, color=DARK):
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=BRAND)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F77B4')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=DARK)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, italic=True, color=GREY)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("\u2139  " + text)
    set_font(run, size=10, italic=True, color=RGBColor(0x1A, 0x5C, 0x8A))
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def add_param(name, desc):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(name + ": ")
    set_font(r1, bold=True, size=10.5)
    r2 = p.add_run(desc)
    set_font(r2, size=10.5)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("MetaBar")
r.font.size  = Pt(32)
r.font.bold  = True
r.font.color.rgb = BRAND

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("User Manual")
r2.font.size  = Pt(18)
r2.font.color.rgb = GREY

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Coskun Lab \u00b7 Georgia Institute of Technology")
r3.font.size  = Pt(12)
r3.font.color.rgb = GREY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Overview")
body("MetaBar is a desktop application for the joint analysis of multiplexed immunofluorescence (IF) and mass spectrometry imaging (MSI) data. It provides a complete, point-and-click workflow \u2014 from raw data loading through image registration, cell segmentation, metabolic quantification, clustering, and graph neural network (GNN) explainability \u2014 with no programming required.")
body("The application runs entirely on your local computer and opens automatically in your web browser when launched.")

heading2("System Requirements")
bullet("Windows 10 or 11 (64-bit)")
bullet("At least 20 GB of free disk space")
bullet("16 GB RAM recommended (32 GB for large datasets)")
bullet("Internet connection required on first launch (to download Python packages)")
bullet("NVIDIA GPU recommended for cell segmentation; CPU fallback is automatic")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 2. INSTALLATION
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Installation")
body("The installer is distributed as two files that must be kept in the same folder:")
bullet("MetaBar_Setup.exe  (run this to install)")
bullet("MetaBar_Setup-1.bin")
note("Both files must be in the same folder before running MetaBar_Setup.exe. Do not move or rename the .bin file.")

heading2("Installation Steps")
numbered("Download both installer files to the same folder on your computer.")
numbered("Double-click MetaBar_Setup.exe.")
numbered("Follow the on-screen wizard. The default installation folder is C:\\Users\\YourName\\MetaBar.")
numbered("Optionally tick 'Create a desktop shortcut' on the last page.")
numbered("Click Finish. The installer will offer to launch the app immediately.")

heading2("First Launch")
body("The first time the app launches, a setup window will appear and automatically download and install the required Python packages (PyTorch, Streamlit, etc.). This takes 10\u201320 minutes and requires an internet connection. It only happens once \u2014 subsequent launches start in a few seconds.")
note("Do not close the setup window while it is running. If it fails, check your internet connection and double-click launch.bat in the installation folder to try again.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 3. LAUNCHING THE APP
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. Launching the App")
body("After installation, launch the app in one of two ways:")
bullet("Double-click the MetaBar shortcut on your desktop.")
bullet("Navigate to the installation folder and double-click launch.bat.")
body("A terminal window will open briefly, then your default web browser will open automatically at http://localhost:8501. The terminal window must stay open while you use the app \u2014 closing it stops the app.")
note("If the browser does not open automatically, open it manually and go to http://localhost:8501.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 4. INTERFACE OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. Interface Overview")
body("The application has two main areas:")
bullet("Sidebar (left panel): data loading and display settings.")
bullet("Main panel (right): image viewer and all analysis steps, organised as collapsible sections.")
body("Analysis sections appear progressively as you complete earlier steps \u2014 you will not see the clustering section, for example, until segmentation and projection have been run.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 5. LOADING DATA
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Loading Data")
body("Data is loaded from the sidebar. Click the IF or MSI tab to switch between the two modalities.")

heading2("5.1  Loading an IF Image")
numbered("Click the IF tab in the sidebar.")
numbered("Click Browse files under 'Multiplex IF TIFF' and select your multi-channel TIFF file.")
numbered("Optionally upload a channel labels file (.txt or .csv, one label per line) under 'Optional IF channel labels'.")
numbered("Click Load IF. A confirmation message shows the number of channels loaded.")

heading2("5.2  Loading MSI Data")
body("MSI data can be loaded in two ways: as a pre-extracted TIFF stack, or directly from raw imzML/IBD files.")

heading3("Option A \u2014 TIFF Stack")
numbered("Click the MSI tab in the sidebar.")
numbered("Select 'TIFF stack' as the input type.")
numbered("Upload your MSI TIFF file and an optional channel labels file.")
numbered("Click Load MSI TIFF.")

heading3("Option B \u2014 imzML / IBD (raw MSI files)")
numbered("Select 'imzML + IBD' as the input type.")
numbered("Upload the .imzML file and the matching .ibd file.")
numbered("Enter the m/z target values you want to extract, either by typing them (comma-separated) or uploading a CSV file with a target_mz column.")
numbered("Set the ppm tolerance (default 5 ppm), output data type, and normalisation method.")
numbered("Click Extract MSI. The app extracts ion images for each target and loads them as a TIFF stack.")
note("Extraction may take several minutes for large imzML files. A progress bar shows the status.")

heading2("5.3  Display Settings")
body("Below the data tabs in the sidebar, the Display section controls how images are rendered:")
add_param("View mode", "Single channels shows one IF and one MSI channel side by side. RGB overlays composites up to three channels into a colour image.")
add_param("Lower / Upper percentile", "Controls contrast stretching. Increase the lower value to suppress background; decrease the upper value to brighten dim signals.")
add_param("Gamma", "Values below 1.0 brighten mid-tones; values above 1.0 darken them.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 6. IMAGE VIEWER
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. Image Viewer")
body("Once both IF and MSI images are loaded, the Viewer section appears in the main panel.")

heading2("Single Channels Mode")
body("Two dropdown menus let you select one IF channel and one MSI channel to display side by side.")

heading2("RGB Overlays Mode")
body("Select up to three IF channels and up to three MSI channels. The first selected channel is mapped to red, the second to green, and the third to blue.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 7. PREPROCESSING
# ══════════════════════════════════════════════════════════════════════════════
heading1("7. Preprocessing")
body("Click 'Preprocessing \u2014 rotate, flip, crop' to expand this section. Preprocessing is optional but recommended when the IF and MSI images have different orientations or contain unwanted border regions.")

heading2("7.1  Rotate and Flip")
add_param("Apply to", "Choose IF, MSI, or Both.")
add_param("Rotation (CCW)", "Rotate by 0\u00b0, 90\u00b0, 180\u00b0, or 270\u00b0 counter-clockwise.")
add_param("Flip", "None, Horizontal, or Vertical.")
body("Click Apply transform to apply the selected rotation and flip.")

heading2("7.2  Crop")
body("Each modality has its own interactive crop panel. To crop:")
numbered("Click and drag on the image to draw a rectangular crop box (shown as an orange dashed rectangle).")
numbered("The pixel coordinates of the selected region are shown below the image.")
numbered("Click Apply crop to [IF/MSI] to apply the crop.")
numbered("Click Clear box to remove the selection without cropping.")
note("Cropping is applied independently to IF and MSI. You can crop one without affecting the other.")

heading2("7.3  Download Processed Stacks")
body("After preprocessing, click Download processed IF TIFF or Download processed MSI TIFF to save the current state of each stack as a multi-channel TIFF file.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 8. REGISTRATION
# ══════════════════════════════════════════════════════════════════════════════
heading1("8. Registration")
body("Registration aligns the MSI image to the IF image using the SIFT feature-matching algorithm in Fiji/ImageJ. Expand the 'Registration \u2014 align MSI to IF via Fiji SIFT' section.")

heading2("8.1  Reference Channels")
add_param("IF reference channel", "Select the IF channel that best represents tissue morphology (e.g. DAPI, Hoechst, or a structural marker).")
add_param("MSI reference channel", "Select the MSI channel that best matches the IF reference (e.g. a lipid with strong tissue contrast).")

heading2("8.2  SIFT Parameters")
body("Expand 'Show / edit SIFT parameters' to adjust the registration algorithm. The defaults work well for most datasets. Key parameters:")
add_param("Expected transformation", "Affine (default) handles rotation, scaling, and shear. Use Rigid or Similarity if you expect only rotation/translation.")
add_param("Maximal alignment error (px)", "Maximum allowed displacement between matched feature pairs. Increase if registration fails on low-contrast images.")
add_param("Inlier ratio", "Minimum fraction of feature matches that must be consistent. Decrease if too few matches are found.")

heading2("8.3  Running Registration")
numbered("Enter the path to your Fiji executable (e.g. C:\\Fiji.app\\ImageJ-win64.exe) in the 'Fiji executable path' field.")
numbered("Click Run Registration.")
numbered("A progress bar tracks the five registration steps. Fiji will open briefly and close automatically \u2014 this is normal.")
numbered("After completion, the aligned MSI stack is available for download.")
note("Registration may take 1\u20135 minutes depending on image size.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 9. ANALYSIS PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
heading1("9. Analysis Pipeline")
body("Expand the 'Analysis Pipeline' section. Before running any analysis step, set the Root output directory \u2014 all results will be saved there in organised subfolders.")
note("Use a dedicated folder per experiment, e.g. Y:\\results\\experiment1. The app will create subfolders automatically.")

heading2("Channel Labels")
body("Edit the IF and MSI channel names in the text boxes (one label per line). These names are used in all downstream outputs, plots, and CSV files. Make sure the order matches the channel order in your TIFF stacks.")

heading2("9.1  Cell Segmentation")
body("Runs nuclear segmentation using the Mesmer deep-learning model.")
add_param("Nuclear channel", "Select the DAPI or Hoechst channel from your IF stack.")
add_param("Tile size (px)", "The image is processed in tiles of this size. 1024 works well for most images; reduce to 512 if you run out of memory.")
add_param("Image resolution (\u00b5m/px)", "The physical pixel size of your IF image. Used by Mesmer to scale the segmentation.")
add_param("Filter area outliers", "Removes cells that are abnormally small or large. Recommended to leave enabled.")
add_param("Area filter method", "mad_log uses median absolute deviation (robust); iqr_log uses interquartile range.")
add_param("DeepCell access token", "Required only on the very first run to download the Mesmer model (~100 MB). Create a free token at https://users.deepcell.org. Leave blank if the model is already cached.")
body("Click Run Segmentation. A progress bar shows tile-by-tile progress. When complete, the number of detected cells is shown and the nuclear mask TIFF can be downloaded.")

heading2("9.2  Nuclei Expansion")
body("Expands the nuclear masks outward to approximate the cytoplasmic compartment.")
add_param("Expansion distance (\u00b5m)", "How far to expand each nucleus. Typical values: 10\u201320 \u00b5m.")
add_param("Pixel size (\u00b5m/px)", "Physical pixel size, used to convert \u00b5m to pixels.")
body("Click Run Expansion. The expanded mask is saved and available for download.")

heading2("9.3  MBP Mask")
body("Generates a binary mask of myelinated tissue regions from the MBP (myelin basic protein) IF channel.")
add_param("MBP channel", "Select the IF channel containing MBP signal.")
add_param("Threshold percentile", "Pixels above this intensity percentile are considered MBP-positive.")
add_param("Gaussian sigma", "Smoothing applied before thresholding. Increase to reduce noise.")
add_param("Min object / hole size (px)", "Small objects and holes below these sizes are removed.")
add_param("Opening / Closing radius", "Morphological operations to clean up the mask boundaries.")
body("Click Run MBP Mask.")

heading2("9.4  LR \u2192 HR Projection")
body("Projects the registered MSI stack from its native (low) resolution to the IF (high) resolution using Gaussian-weighted interpolation, then computes per-cell mean MSI intensities.")
add_param("Gaussian sigma (LR pixels)", "Controls the smoothness of the interpolation. Default 0.75 is appropriate for most datasets.")
add_param("Gaussian radius (LR pixels)", "Neighbourhood size for interpolation. Default 2.")
body("Click Run Projection. Two CSV files are produced: one with per-cell intensities within nuclear masks, and one within expanded nuclear masks. Both can be downloaded.")

heading2("9.5  Superpixel Segmentation")
body("Partitions the tissue into superpixels using the SLIC algorithm, then computes per-superpixel mean MSI intensities. This step is optional \u2014 clustering and GNN will run on cells only if superpixels are skipped.")
add_param("MBP channel for SLIC", "The IF channel used to guide superpixel boundaries.")
add_param("Target segments", "Approximate number of superpixels. More segments = finer spatial resolution.")
add_param("Compactness", "Higher values produce more square superpixels; lower values follow image boundaries more closely.")
add_param("Downsample factor", "Reduces image size before SLIC for speed. The result is upsampled back afterwards.")
body("Click Run Superpixel Segmentation. A CSV of per-superpixel MSI intensities is produced.")

heading2("9.6  Clustering")
body("Clusters cells (and optionally superpixels) based on their MSI intensity profiles using PCA, UMAP, Leiden community detection, and k-means. Superpixel segmentation is not required to run clustering.")
add_param("Channels to use", "Select which MSI channels to include as features. Deselect channels that are noisy or irrelevant.")
add_param("PCA components", "Number of principal components to compute before UMAP. Default 10.")
add_param("UMAP neighbors", "Controls the local neighbourhood size in UMAP. Higher values give a more global view.")
add_param("UMAP min dist", "Minimum distance between points in the UMAP embedding. Lower values produce tighter clusters.")
add_param("Leiden resolutions", "Comma-separated list of resolution values. Higher resolution = more, smaller clusters.")
add_param("KMeans k values", "Comma-separated list of k values to try.")
add_param("Row z-score", "Normalises each cluster's mean intensity by z-score in the matrix plot.")
body("Click Run Clustering. UMAP plots, matrix plots, and spatially coloured cluster maps are saved for cells and superpixels (if available).")

heading2("9.7  Custom Data")
body("The Custom Data tab lets you bypass the segmentation pipeline entirely by uploading your own cell mask and/or annotation file.")

heading3("Uploading a Custom Mask")
body("Upload a 2-D integer label TIFF where each pixel value is the cell ID (0 = background). Optionally enable 'Expand mask before quantification' to expand each cell outward before computing per-cell MSI intensities \u2014 set the expansion distance and pixel size, then click Save mask & run projection.")
note("After saving, the app immediately runs MSI projection and cell quantification on your mask. You can then proceed directly to Clustering or GNN Explainability.")

heading3("Uploading Cell Annotations / Phenotypes")
body("Upload a single-column CSV or TXT file where row N is the label for cell ID N (row 1 = cell 1, row 2 = cell 2, etc.). No header is needed. The annotation is saved to the annotations/cells/ folder and appears in the napari viewer under 'Cell Annotations', and is available as a GNN prediction target in the Annotations GNN tab.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 10. POSITIVITY THRESHOLDING
# ══════════════════════════════════════════════════════════════════════════════
heading1("10. Positivity Thresholding")
body("Expand the 'Positivity Thresholding' section. This step assigns a binary positive/negative label to each cell for each selected MSI channel, based on the distribution of per-cell intensities.")
add_param("Threshold method", "top_component_quantile (recommended): fits a Gaussian mixture model and uses a quantile of the highest-intensity component as the threshold. gmm: uses the intersection of two GMM components. otsu: Otsu's method. upper_quantile: simple percentile threshold.")
add_param("GMM components", "Number of components in the mixture model (for top_component_quantile).")
add_param("Component quantile", "Quantile within the top GMM component used as the threshold. 0.5 = median of the top component.")
add_param("Fallback quantile", "Used when the GMM fails or produces extreme class imbalance.")
add_param("Min / Max positive fraction", "Clamps the fraction of positive cells to a reasonable range.")
add_param("Save positivity overlay TIFFs", "Saves a colour-coded TIFF per channel (red = positive, grey = negative) for visual inspection.")
body("Select the MSI channels to threshold and click Run Positivity Thresholding. A results table shows the threshold, method used, and positive/negative cell counts for each channel.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 11. GNN EXPLAINABILITY
# ══════════════════════════════════════════════════════════════════════════════
heading1("11. GNN Explainability")
body("Expand the 'GNN Explainability' section. This step trains graph neural networks on the spatial cell graph to identify which MSI channels are most predictive of cell identity or metabolic state.")
body("Set the Results root folder to the same folder used in the Analysis Pipeline.")

heading2("11.1  Positivity GNN (Binary)")
body("Predicts whether each cell is positive for a selected IF marker, using the binary labels from the Positivity Thresholding step.")
add_param("Markers", "Select one or more markers to train. Each marker is trained as a separate binary classification task.")
add_param("Feature channels", "MSI channels used as node features. Consider excluding the channel that defines the positivity label to avoid data leakage.")
add_param("Radius (\u00b5m)", "Cells within this distance are connected by an edge in the spatial graph.")
add_param("K-folds", "Number of cross-validation folds.")
add_param("Model", "GraphSAGE (default), GCN, or GATv2.")
add_param("Hidden dim / Layers", "Network architecture. Larger values increase capacity but require more memory.")
add_param("Max epochs / Patience", "Training stops after this many epochs or when validation loss stops improving.")
add_param("Explainability method", "saliency (gradient-based) or occlusion (feature masking).")
add_param("Top-k features", "Number of most important MSI channels to report.")
body("Click Run Binary GNN. Results show ROC-AUC and F1 scores per marker, and a feature importance bar chart.")

heading2("11.2  Clustering GNN (Multiclass)")
body("Predicts cluster membership for a selected Leiden or k-means clustering result.")
add_param("Clustering result", "Select which clustering column to use as the prediction target.")
body("All other parameters are the same as the binary GNN. Click Run Multiclass GNN. Per-class feature importance charts are shown for each cluster.")

heading2("11.3  Annotations GNN (Multiclass)")
body("Predicts cell annotation / phenotype labels uploaded via the Custom Data tab. The annotation column must exist in the annotations/cells/ folder.")
add_param("Annotation column", "Select which annotation column to use as the prediction target.")
body("All other parameters are the same as the clustering GNN. Click Run Annotation GNN. Per-class feature importance charts are shown for each annotation class.")
note("If no annotation columns appear, upload an annotation file in the Custom Data tab first.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 12. CROSS-SAMPLE COMPARATIVE ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
heading1("12. Cross-Sample Comparative Analysis")
body("Expand the 'Cross-Sample Comparative Analysis' section. This step compares GNN feature importance results across multiple samples or conditions.")
numbered("Click Add to register each sample: enter a short name (e.g. control, treated) and the path to that sample's results root folder.")
numbered("Add at least two samples. A green tick indicates the folder was found; a red cross means the path is invalid.")
numbered("Set the Comparison output folder.")
numbered("Set Top-N features per plot.")

body("Two comparison modes are available as tabs:")
add_param("Positivity (binary)", "Compares per-marker binary GNN feature importance from gnn_explainability/binary/ across samples.")
add_param("Annotations (multiclass)", "Compares per-class annotation GNN feature importance from gnn_explainability/annotations/ across samples. Run the Annotations GNN tab for each sample first.")
body("Click the Run button for the desired mode. The app generates grouped bar charts and heatmaps comparing the relative importance of each MSI channel across conditions.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 13. INTERACTIVE VIEWER (napari)
# ══════════════════════════════════════════════════════════════════════════════
heading1("13. Interactive Viewer (napari)")
body("Expand the 'Interactive Viewer (napari)' section. This launches a separate napari window for interactive spatial exploration of all analysis results.")
body("The viewer loads:")
bullet("All IF fluorescence channels")
bullet("The Gaussian-projected MSI stack at IF resolution")
bullet("Nuclear and expanded nuclear masks")
bullet("Superpixel label mask (if available)")
bullet("Cluster-coloured overlays for all Leiden and k-means results")
bullet("Cell annotation overlays (if annotations have been uploaded via Custom Data)")
body("Set the Results root folder (same as the Analysis Pipeline) and click Launch napari. The napari window opens as a separate application. You can inspect individual cells by hovering or clicking on the image.")
body("In the Cluster Overlays panel on the right, three groups are available: Cells (clustering results), Superpixels (if available), and Annotations (custom phenotype labels). Select the method and value from each group's dropdowns to paint the corresponding overlay.")
note("The MSI projection TIFF (projected_stack_all_channels__full_hr__gaussian.tif) must be present in the projection/ subfolder for the MALDI overlay layers to load. If this file was not included in downloaded results, run the LR \u2192 HR Projection step first using the raw demo data.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 14. OUTPUT FILES
# ══════════════════════════════════════════════════════════════════════════════
heading1("14. Output Files")
body("All outputs are saved in the Results root folder you specified, organised as follows:")

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.8)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("results/\n")
set_font(r, bold=True)
lines = [
    "  segmentation/          nuclear_mask.tif, nuclear_mask_expanded.tif",
    "  projection/            projected_stack_*.tif, cell_level_metabolic_table_*.csv",
    "  superpixels/           mbp_superpixels_label_mask.tif, *_mean_intensity_matrix.csv",
    "  clustering/cells/      cells__clustered.csv, UMAP plots, matrix plots, coloured masks",
    "  clustering/superpixels/ superpixels__clustered.csv, UMAP plots, matrix plots",
    "  annotations/cells/     cells__clustered.csv, {annotation}__colors.csv, colored masks",
    "  positivity/            protein_marker_thresholds.csv, cell_binary_labels.csv, overlay TIFFs",
    "  gnn_explainability/    feature_importance_topk.png, per-fold metrics",
]
for line in lines:
    rr = p.add_run(line + "\n")
    set_font(rr, size=10)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 15. TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════
heading1("15. Troubleshooting")

heading2("App does not open after double-clicking the shortcut")
body("Open a Command Prompt, navigate to the installation folder, and run launch.bat directly to see the error message:")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
r = p.add_run('cd "C:\\Users\\YourName\\MetaBar"\nlaunch.bat')
r.font.name = 'Courier New'
set_font(r, size=10)

heading2("First-run setup fails with a network error")
body("Check your internet connection and try again by deleting the file envs\\.setup_complete in the installation folder, then running launch.bat again.")

heading2("napari window does not open")
body("If the napari viewer fails to launch, expand the 'napari log' section in the app to see the error. Common causes:")
bullet("napari is not installed: run fix_napari.bat from the MetaBar installation folder.")
bullet("Qt initialisation error: restart the app and try again.")
bullet("Missing projection TIFF: run LR \u2192 HR Projection first if the gaussian TIFF is not present.")

heading2("Cell segmentation fails or produces no cells")
bullet("Check that you selected the correct nuclear channel (DAPI/Hoechst).")
bullet("Try reducing the tile size to 512 if you have limited GPU memory.")
bullet("If running on CPU, segmentation will be slow (30\u201360 min for large images) \u2014 this is normal.")
bullet("On first run, a DeepCell access token is required to download the model. Create one free at https://users.deepcell.org.")

heading2("Registration produces a misaligned result")
bullet("Try a different pair of reference channels with stronger tissue contrast.")
bullet("Increase the Maximal alignment error parameter.")
bullet("Ensure both images cover the same tissue region after preprocessing.")

heading2("Out of memory errors during analysis")
bullet("Reduce the tile size in segmentation.")
bullet("Increase the downsample factor in superpixel segmentation.")
bullet("Close other applications to free RAM.")

heading2("GNN training is very slow")
body("GNN training uses the GPU if available. If no GPU is detected, training runs on CPU and may take significantly longer. Reduce Max epochs or K-folds to speed up a test run.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 16. CONTACT
# ══════════════════════════════════════════════════════════════════════════════
heading1("16. Contact and Support")
body("For questions, bug reports, or feature requests, please open an issue on the project GitHub repository:")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.8)
r = p.add_run("https://github.com/coskunlab/MetaBar")
r.font.name = 'Courier New'
set_font(r, size=10.5, color=BRAND)

body("Coskun Lab \u00b7 Georgia Institute of Technology")

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r'Y:\coskun-lab\Efe\Metabobarcoding\app\MetaBar_User_Manual.docx'
doc.save(out_path)
print(f"Saved: {out_path}")
